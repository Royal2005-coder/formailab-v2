#!/usr/bin/env python3
"""Generate the Vietnamese AILAB Survey technical/business handover DOCX."""

from __future__ import annotations

import json
import os
import subprocess
from datetime import datetime, timezone
from pathlib import Path
from xml.sax.saxutils import escape

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs" / "handover"
ASSET_DIR = OUT_DIR / "assets"
OUT_FILE = OUT_DIR / "AILAB_Survey_Tai_lieu_ban_giao_ky_thuat_nghiep_vu_2026.docx"
BLUE = "2367D1"
NAVY = "10233F"
TEAL = "00B8A9"
LIGHT = "EAF2FF"
GRAY = "667085"
GREEN = "067647"
AMBER = "B54708"
RED = "B42318"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def keep_with_next(paragraph) -> None:
    paragraph.paragraph_format.keep_with_next = True


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Trang ")
    run.font.size = Pt(8)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)


def add_toc(paragraph) -> None:
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "Nhấn Ctrl+A rồi F9 trong Microsoft Word để cập nhật mục lục."
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for node in (fld_char, instr, separate, text, end):
        run._r.append(node)


def add_hyperlink(paragraph, text: str, url: str):
    part = paragraph.part
    rel_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    props = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    props.extend((color, underline))
    run.append(props)
    node = OxmlElement("w:t")
    node.text = text
    run.append(node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths=None, font_size=8):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for idx, label in enumerate(headers):
        cell = hdr.cells[idx]
        set_cell_shading(cell, NAVY)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        r = p.add_run(label)
        r.bold = True
        r.font.color.rgb = RGBColor(255, 255, 255)
        r.font.size = Pt(font_size)
    for row_idx, row in enumerate(rows):
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            if row_idx % 2:
                set_cell_shading(cells[idx], "F8FAFC")
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            r = p.add_run(str(value))
            r.font.size = Pt(font_size)
        if widths:
            for idx, width in enumerate(widths):
                cells[idx].width = Cm(width)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_bullets(doc: Document, items: list[str], level=0) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
        p.add_run(item)


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_callout(doc: Document, title: str, body: str, kind="info") -> None:
    colors = {"info": (LIGHT, BLUE), "ok": ("ECFDF3", GREEN), "warn": ("FFFAEB", AMBER), "risk": ("FEF3F2", RED)}
    fill, color = colors[kind]
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    r = p.add_run(title + " — ")
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(color)
    p.add_run(body)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def svg_box_diagram(name: str, title: str, columns: list[list[tuple[str, str]]], arrows: list[tuple[int, int, str]] | None = None) -> Path:
    width, height = 1400, 180 + max(len(c) for c in columns) * 150
    col_width = width / len(columns)
    boxes = []
    positions = {}
    for ci, col in enumerate(columns):
        for ri, (key, label) in enumerate(col):
            x = int(ci * col_width + 45)
            y = 110 + ri * 150
            w = int(col_width - 90)
            h = 92
            positions[key] = (x, y, w, h)
            fill = ["#EAF2FF", "#E8FAF7", "#FFF6E5", "#F3EEFF"][ci % 4]
            boxes.append(f'<rect x="{x}" y="{y}" width="{w}" height="{h}" rx="16" fill="{fill}" stroke="#2367D1" stroke-width="2"/>')
            lines = label.split("\n")
            for li, line in enumerate(lines):
                boxes.append(f'<text x="{x+w/2}" y="{y+38+li*24}" text-anchor="middle" font-family="Arial" font-size="20" fill="#10233F">{escape(line)}</text>')
    connectors = []
    if arrows:
        for src_idx, dst_idx, label in arrows:
            keys = list(positions)
            if src_idx >= len(keys) or dst_idx >= len(keys):
                continue
            sx, sy, sw, sh = positions[keys[src_idx]]
            dx, dy, dw, dh = positions[keys[dst_idx]]
            x1, y1 = sx + sw, sy + sh / 2
            x2, y2 = dx, dy + dh / 2
            if dx < sx:
                x1, x2 = sx, dx + dw
            connectors.append(f'<path d="M{x1},{y1} C{(x1+x2)/2},{y1} {(x1+x2)/2},{y2} {x2},{y2}" fill="none" stroke="#667085" stroke-width="2" marker-end="url(#arrow)"/>')
            if label:
                connectors.append(f'<text x="{(x1+x2)/2}" y="{(y1+y2)/2-8}" text-anchor="middle" font-family="Arial" font-size="15" fill="#475467">{escape(label)}</text>')
    svg = f'''<svg xmlns="http://www.w3.org/2000/svg" width="{width}" height="{height}">
<defs><marker id="arrow" markerWidth="10" markerHeight="7" refX="9" refY="3.5" orient="auto"><polygon points="0 0,10 3.5,0 7" fill="#667085"/></marker></defs>
<rect width="100%" height="100%" fill="white"/><text x="700" y="55" text-anchor="middle" font-family="Arial" font-size="30" font-weight="700" fill="#10233F">{escape(title)}</text>
{''.join(connectors)}{''.join(boxes)}</svg>'''
    svg_path = ASSET_DIR / f"{name}.svg"
    png_path = ASSET_DIR / f"{name}.png"
    svg_path.write_text(svg, encoding="utf-8")
    js = "const sharp=require('sharp'); sharp(process.argv[1]).png().toFile(process.argv[2]).catch(e=>{console.error(e);process.exit(1)})"
    subprocess.run(["node", "-e", js, str(svg_path), str(png_path)], cwd=ROOT, check=True)
    return png_path


def add_figure(doc: Document, path: Path, caption: str, width=6.8) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(width))
    cap = doc.add_paragraph(caption)
    cap.style = "Caption"
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER


def source(doc: Document, paths: str) -> None:
    p = doc.add_paragraph()
    p.style = "Source Note"
    p.add_run("Nguồn kiểm chứng: ").bold = True
    p.add_run(paths)


def page_break(doc: Document) -> None:
    doc.add_page_break()


def build_diagrams() -> dict[str, Path]:
    diagrams = {}
    diagrams["context"] = svg_box_diagram(
        "01-system-context", "Bối cảnh hệ thống AILAB Survey",
        [[("actors", "Người trả lời\nBA · PM · Admin")], [("web", "AILAB Survey Web\nNext.js App Router"), ("survey", "Survey Runtime\nAdaptive engine")], [("data", "PostgreSQL · Valkey\nObject Storage"), ("ai", "Hub · Worker · Taxonomy\nGemini AI")], [("external", "SMTP · SSO/OIDC/SAML\nAPI/Webhook/Cube")]],
        [(0, 1, "HTTPS"), (1, 3, "responses"), (1, 4, "API"), (4, 5, "model")],
    )
    diagrams["components"] = svg_box_diagram(
        "02-component-architecture", "Kiến trúc component và module",
        [[("routes", "apps/web/app\nPages · Route Handlers"), ("ui", "apps/web/modules\nFeature UI")], [("actions", "Server Actions\nAccess checks"), ("services", "apps/web/lib\nBusiness services")], [("packages", "packages/*\nAI · DB · Survey compiler"), ("prisma", "Prisma ORM\nSchema & migrations")], [("infra", "Docker · Caddy\nCube · Hub workers")]],
        [(0, 2, "calls"), (1, 2, "events"), (2, 3, "service"), (3, 4, "shared"), (4, 5, "queries"), (5, 6, "deploy")],
    )
    diagrams["deployment"] = svg_box_diagram(
        "03-deployment-topology", "Deployment topology production",
        [[("client", "Browser / Mobile\nHTTPS")], [("edge", "Cloudflare / Caddy\nTLS · reverse proxy")], [("next", "formbricks-ai-lab-staging\nNext.js + workers"), ("cube", "Cube analytics")], [("pg", "PostgreSQL"), ("redis", "Valkey/Redis"), ("hub", "Hub API + Worker\nTaxonomy service"), ("storage", "RustFS/S3")]],
        [(0, 1, "443"), (1, 2, "HTTP"), (2, 4, "SQL"), (2, 5, "cache"), (2, 6, "feedback"), (2, 7, "files")],
    )
    diagrams["ai"] = svg_box_diagram(
        "04-ai-native-pipeline", "AI-native feedback pipeline",
        [[("ingest", "Nguồn phản hồi\nSurvey · CSV · API")], [("normalize", "Chuẩn hóa & mapping\nFeedbackDirectory tenant")], [("embed", "Embedding 768D\nGemini embedding-001"), ("sentiment", "Sentiment · Emotion\nGemini 2.5 Flash")], [("taxonomy", "Clustering / taxonomy\nTopic · Subtopic"), ("search", "Semantic search\nDashboard & chart")]],
        [(0, 1, "ingest"), (1, 2, "async"), (1, 3, "async"), (2, 4, "vectors"), (3, 4, "labels"), (4, 5, "query")],
    )
    diagrams["import"] = svg_box_diagram(
        "05-import-bpmn", "Luồng BPMN rút gọn: CSV/XLSX → Survey",
        [[("file", "Upload CSV/XLSX\n≤ 10 MiB")], [("parse", "Decode UTF-8 / Win-1258\nParse workbook sheets"), ("validate", "Canonical validation\nDiagnostics")], [("compile", "Compile blocks\nLogic · variables · quotas"), ("review", "Preview & approve")], [("commit", "Transaction commit\nRegistry + version + draft"), ("publish", "Publish / rollback\nAudit & lifecycle")]],
        [(0, 1, "parse"), (1, 2, "schema"), (2, 3, "valid"), (3, 4, "preview"), (4, 5, "commit"), (5, 6, "release")],
    )
    diagrams["adaptive"] = svg_box_diagram(
        "06-adaptive-sequence", "Sequence: trả lời khảo sát thích ứng",
        [[("respondent", "Người trả lời")], [("runtime", "Survey Runtime\nReact/UMD bundle")], [("compiler", "Compiled blocks\nExpression evaluator"), ("api", "Response API\nvalidation")], [("db", "PostgreSQL\nResponse · Display"), ("pipeline", "Response pipeline\nIntegrations & AI")]],
        [(0, 1, "answer"), (1, 2, "evaluate"), (2, 1, "next block"), (1, 3, "submit"), (3, 4, "transaction"), (4, 5, "job")],
    )
    diagrams["rbac"] = svg_box_diagram(
        "07-rbac-usecase", "RBAC và phạm vi đa tenant",
        [[("user", "User / Account\n2FA · SSO")], [("membership", "Organization Membership\nowner · manager · member"), ("team", "TeamUser\nadmin · contributor")], [("workspace", "WorkspaceTeamPermission\nread · readWrite · manage"), ("apikey", "API key scopes\norganizationAccess")], [("resources", "Survey · Contact · Segment\nDashboard · Dataset")]],
        [(0, 1, "member"), (0, 2, "team"), (1, 3, "org scope"), (2, 3, "workspace"), (3, 5, "authorize"), (4, 5, "authorize")],
    )
    diagrams["erd"] = svg_box_diagram(
        "08-core-erd", "ERD logic: các thực thể cốt lõi",
        [[("org", "Organization\nMembership · Invite"), ("user", "User · Account · Session")], [("workspace", "Workspace\nTeam · WorkspaceTeam"), ("survey", "Survey\nRegistry · Version · Publication")], [("response", "Response · Display\nQuota links · Tags"), ("contact", "Contact · Attributes\nSegment")], [("feedback", "FeedbackDirectory\nSource · Field mapping"), ("analysis", "Chart · Dashboard\nDashboardWidget")]],
        [(0, 2, "1:N"), (1, 0, "N:M"), (2, 3, "1:N"), (3, 4, "1:N"), (5, 4, "0:N"), (2, 6, "N:M"), (2, 7, "1:N")],
    )
    return diagrams


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.6)
    section.left_margin = Cm(2.1)
    section.right_margin = Cm(1.8)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(9.5)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.12
    for idx, size in ((1, 18), (2, 14), (3, 11)):
        style = styles[f"Heading {idx}"]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(NAVY if idx == 1 else BLUE)
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.keep_with_next = True
    styles["Title"].font.name = "Arial"
    styles["Title"].font.color.rgb = RGBColor.from_string(NAVY)
    styles["Caption"].font.name = "Arial"
    styles["Caption"].font.size = Pt(8)
    styles["Caption"].font.italic = True
    if "Source Note" not in [s.name for s in styles]:
        style = styles.add_style("Source Note", WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = "Arial"
        style.font.size = Pt(7.5)
        style.font.color.rgb = RGBColor.from_string(GRAY)
        style.font.italic = True
    for sec in doc.sections:
        header = sec.header.paragraphs[0]
        header.text = "AILAB SURVEY  •  TÀI LIỆU BÀN GIAO  •  04/08/2026"
        header.style = styles["Source Note"]
        footer = sec.footer.paragraphs[0]
        footer.text = "Tài liệu kỹ thuật & nghiệp vụ — phiên bản 1.0  |  Nội bộ"
        add_page_number(footer)
    settings = doc.settings.element
    update_fields = OxmlElement("w:updateFields")
    update_fields.set(qn("w:val"), "true")
    settings.append(update_fields)


def heading(doc: Document, text: str, level=1) -> None:
    doc.add_heading(text, level=level)


def build_document() -> Path:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    diagrams = build_diagrams()
    scan = json.loads((ROOT / ".ua/intermediate/scan-result.json").read_text())
    knowledge_path = ROOT / ".ua/knowledge-graph.json"
    knowledge = json.loads(knowledge_path.read_text()) if knowledge_path.exists() else {"nodes": [], "edges": []}
    demo_password = os.environ.get("AILAB_DEMO_PASSWORD", "[ĐƯỢC QUẢN LÝ NGOÀI GIT]")
    doc = Document()
    configure_document(doc)

    # Cover
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    logo = ASSET_DIR / "brand-logo.png"
    if logo.exists():
        p.add_run().add_picture(str(logo), width=Inches(2.7))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("AILAB SURVEY")
    r.bold = True; r.font.size = Pt(30); r.font.color.rgb = RGBColor.from_string(NAVY)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("TÀI LIỆU BÀN GIAO KỸ THUẬT, NGHIỆP VỤ\nVÀ HƯỚNG DẪN VẬN HÀNH TOÀN HỆ THỐNG")
    r.bold = True; r.font.size = Pt(17); r.font.color.rgb = RGBColor.from_string(BLUE)
    doc.add_paragraph()
    meta = add_table(doc, ["Thuộc tính", "Giá trị"], [
        ["Phiên bản tài liệu", "1.0 — Production handover baseline"],
        ["Ngày chốt bằng chứng", "04/08/2026 (UTC)"],
        ["Production", "https://formailab.royalai.dev"],
        ["Workspace tham chiếu", "cms7eskq6000901o9ul1ipheu"],
        ["Commit khảo sát", "ef92d78aa86916e6fbf4bc8a43044a4a4328154d"],
        ["Đối tượng", "BA · PM · PO · Stakeholder · Dev · DevOps · QA · End user"],
        ["Phân loại", "Nội bộ — không chứa khóa, token hoặc mật khẩu"],
    ], widths=[5, 11], font_size=9)
    page_break(doc)

    heading(doc, "Lịch sử và kiểm soát tài liệu")
    add_table(doc, ["Phiên bản", "Ngày", "Nội dung", "Trạng thái"], [["1.0", "04/08/2026", "Bản bàn giao tổng hợp codebase, production evidence, nghiệp vụ và vận hành", "Baseline"]])
    add_callout(doc, "Nguyên tắc xác thực", "Tài liệu phân biệt rõ mã nguồn có sẵn, chức năng đã chạy production, dữ liệu demo/canary và rủi ro chưa đóng. Không coi mock data là bằng chứng AI-native.", "info")
    heading(doc, "Mục lục")
    add_toc(doc.add_paragraph())
    page_break(doc)

    heading(doc, "1. Tóm tắt điều hành")
    doc.add_paragraph("AILAB Survey là nền tảng quản trị khảo sát đa tenant được mở rộng từ Formbricks, kết hợp trình tạo khảo sát, nhập ngân hàng câu hỏi CSV/XLSX, logic rẽ nhánh thích ứng, danh bạ/segment, Unify Feedback, dashboard/chart và pipeline AI-native. Web app sử dụng Next.js App Router; dữ liệu nghiệp vụ nằm trong PostgreSQL qua Prisma; Redis/Valkey hỗ trợ cache/job; Hub/worker/taxonomy và Gemini đảm nhiệm embedding, sentiment/emotion, semantic search và phân loại chủ đề.")
    add_figure(doc, diagrams["context"], "Hình 1 — Bối cảnh hệ thống và các tích hợp chính")
    heading(doc, "1.1 Kết luận production-readiness tại thời điểm bàn giao", 2)
    add_table(doc, ["Miền", "Kết luận", "Bằng chứng", "Mức tin cậy"], [
        ["Nền tảng web", "Đang hoạt động", "/health=200; /api/v2/health xác nhận main database và cache", "Cao"],
        ["Importer adaptive", "Mã nguồn và E2E có transaction/cleanup; template XLSX tải được", "Playwright + service tests + schema registry", "Cao"],
        ["AI embedding/search", "AI thật, lưu dữ liệu production, không phải mock", "Canary được embedding 768D; semantic search trả score; isolation/duplicate/auth đã kiểm", "Cao"],
        ["Sentiment/emotion", "Pipeline thật hoạt động nhưng độ phủ phụ thuộc dữ liệu/model", "280 bản ghi text có sentiment; 171 có emotion tại lần kiểm", "Trung bình-cao"],
        ["Taxonomy", "Clustering thật; nhãn hiện tại có bước curated", "815 embeddings, 13 clusters, 36 nodes; model từng tạo nhãn lỗi", "Trung bình"],
        ["AI chart", "Lỗi schema filter đã sửa và deploy", "Regression 13/13 unit tests; image mới health=200", "Cao"],
        ["Việt hóa", "Phần lớn app shell đã Việt hóa, chưa đạt 100%", "Production snapshot còn Required/Please fill out this field/Rate...", "Cao"],
        ["SSO/RBAC", "Đầy đủ cấu trúc và enforcement trong code; cần UAT IdP thực tế", "Better Auth, SAML Jackson, OIDC, Membership/Team/WorkspaceTeam", "Trung bình-cao"],
    ])
    add_callout(doc, "Quyết định go-live", "Có thể vận hành có kiểm soát cho khảo sát/import/dashboard và AI feedback. Chưa nên tuyên bố 'hoàn thiện 100%' cho i18n, taxonomy label automation và toàn bộ SSO nếu chưa đóng các mục P0/P1 ở Chương 18.", "warn")

    heading(doc, "2. Phạm vi, phương pháp và bằng chứng")
    doc.add_paragraph(f"Knowledge graph được dựng từ {scan['totalFiles']:,} tệp sản phẩm và {sum(len(v) for v in scan['importMap'].values()):,} cạnh import; đã loại mã tooling `.agents` ngoại trừ skill form-builder. Kết quả cuối có {len(knowledge['nodes']):,} node và {len(knowledge['edges']):,} cạnh, được phân tích qua 197 cộng đồng mã nguồn. Phạm vi bao gồm apps/web, packages, database migrations, docs, docker, Playwright, Testbank và form-builder skill.")
    add_table(doc, ["Lớp kiểm chứng", "Cách thực hiện", "Đầu ra"], [
        ["Static code", "Scanner + import map + structural extraction + đọc sâu module trọng yếu", ".ua/knowledge graph và trích dẫn đường dẫn"],
        ["Database", "Đọc Prisma schema/migrations; truy vấn read-only production", "ERD logic, tenant scope, inventory dữ liệu"],
        ["Runtime", "Docker/container health, logs, API behavior", "Trạng thái deployment và rủi ro"],
        ["Browser", "Playwright Desktop Chrome, trace/video/screenshot", "JUnit/JSON và ảnh production"],
        ["AI-native", "Canary thật, embedding/model output, semantic search, tenant isolation", "Bằng chứng không dùng mock"],
        ["Tài liệu", "Sinh DOCX, mở lại bằng python-docx, kiểm table/image/heading", "Artifact bàn giao kiểm định"]
    ])
    source(doc, ".ua/intermediate/scan-result.json; playwright.production.config.ts; test-results/production; packages/database/schema.prisma")

    heading(doc, "3. Kiến trúc tổng thể")
    add_figure(doc, diagrams["components"], "Hình 2 — Phân lớp route/UI/action/service/package/data/infra")
    heading(doc, "3.1 Monorepo và trách nhiệm", 2)
    add_table(doc, ["Khu vực", "Trách nhiệm", "Ví dụ"], [
        ["apps/web/app", "Next.js App Router: page, layout, route handler, middleware", "workspaces/[workspaceId], api/v3, auth"],
        ["apps/web/modules", "Feature UI + action/lib theo miền", "ai-lab-survey, contacts, analysis, unify-feedback, sso"],
        ["apps/web/lib", "Business services, access, cache, server utilities", "survey, workspace, organization, i18n"],
        ["packages/database", "Prisma client, schema, migrations, seed", "Survey, Response, AI registry, RBAC"],
        ["packages/survey-compiler", "Biên dịch cấu trúc survey/import thành runtime contract", "expression, logic, blocks"],
        ["packages/surveys + survey-ui", "Survey respondent runtime và UI bundle", "UMD/ESM, adaptive rendering"],
        ["packages/ai", "Provider abstraction, model calls và structured output", "Gemini/OpenAI-compatible"],
        ["docker / docs / Testbank", "Deployment, operator docs, fixtures và test bank", "Caddy, Cube, templates CSV/XLSX"]
    ])
    heading(doc, "3.2 Nguyên tắc Next.js áp dụng", 2)
    add_bullets(doc, [
        "Route groups như `(app)` và `(analysis)` tổ chức layout nhưng không xuất hiện trong URL.",
        "Server Components là mặc định; Client Components chỉ dùng khi cần state, effect hoặc browser API.",
        "Server Actions bọc service calls và trả `{data}`/`{error}`; Route Handlers phục vụ API contract.",
        "Cache request-level dùng React `cache()`; cache chia sẻ dùng cache.withCache/Redis và createCacheKey.",
        "Authorization phải kiểm tra ở data-access/service boundary, không chỉ ẩn nút trên UI."
    ])
    source(doc, "node_modules/next/dist/docs/01-app/01-getting-started/02-project-structure.md; AGENTS.md; apps/web/app; apps/web/modules")

    heading(doc, "4. Hạ tầng và triển khai")
    add_figure(doc, diagrams["deployment"], "Hình 3 — Topology container và luồng production")
    add_table(doc, ["Thành phần", "Vai trò", "Dữ liệu/state", "Health/quan sát"], [
        ["ai-lab-edge / Caddy", "TLS và reverse proxy", "Stateless", "HTTP status, certificate"],
        ["formbricks-ai-lab-staging", "Next.js web/API và tác vụ nền", "Stateless ngoài cache/tmp", "/health, /api/v2/health"],
        ["formbricks-postgres", "OLTP: survey, users, RBAC, dashboard", "Persistent volume", "SQL health/migrations"],
        ["formbricks-valkey", "Cache, queue coordination", "Persistent tùy cấu hình", "PING/connection"],
        ["formbricks-hub", "Feedback records/semantic API", "Hub database", "/health, 401 auth guard"],
        ["formbricks-hub-worker", "Embedding/sentiment/emotion async", "Queue + database", "job lag, retry, DLQ"],
        ["formbricks-taxonomy", "Clustering và taxonomy runs", "Run state/nodes", "run status/counts"],
        ["formbricks-cube", "Semantic analytics layer", "Cube cache", "query latency/errors"],
        ["RustFS/S3", "Upload files/assets", "Object storage", "bucket/permission/retention"]
    ])
    heading(doc, "4.1 Build, runtime và version", 2)
    add_table(doc, ["Thành phần", "Version/baseline"], [
        ["Node.js", ".nvmrc = 24.14.0; package engines hỗ trợ 20.19/22.12/24"],
        ["pnpm / Turbo", "pnpm 11.7.0; turbo 2.9.14"],
        ["Next.js / React", "Next 16.2.11; React 19.2.6"],
        ["Prisma / PostgreSQL", "Prisma 7.8.0; PostgreSQL container"],
        ["Playwright / Vitest", "Playwright 1.58.2; Vitest 4.1.6"],
        ["Tailwind / Recharts", "Tailwind 4.3.1; Recharts 2.15.3"],
    ])
    add_callout(doc, "Node single source of truth", "CI/dev/prod phải khóa theo `.nvmrc`; các script production nên chạy bằng Node 24 để tránh sai khác runtime. Survey packages cần build `--force` và hard refresh do UMD bundle/cache.", "warn")
    source(doc, ".nvmrc; package.json; apps/web/package.json; docker/ai-lab/*; docker/docker-compose.yml")

    heading(doc, "5. Mô hình dữ liệu và đa tenant")
    add_figure(doc, diagrams["erd"], "Hình 4 — ERD logic rút gọn; chi tiết trường nằm trong Prisma schema")
    add_table(doc, ["Cụm", "Thực thể", "Ràng buộc nghiệp vụ"], [
        ["Identity", "User, Account, Session, TwoFactor, VerificationToken", "Một user có nhiều identity/session; token không đưa vào log/doc"],
        ["Organization", "Organization, Membership, Invite, OrganizationBilling", "Role owner/manager/member; invite có expiry và creator/acceptor"],
        ["Workspace/Team", "Workspace, Team, TeamUser, WorkspaceTeam", "Permission read/readWrite/manage; mọi resource phải scope workspace/org"],
        ["Survey", "Survey, AILabSurveyRegistry, Version, Publication, ImportJob", "Version hóa, lifecycle, publish/rollback, diagnostics"],
        ["Response", "Response, Display, Quota, ResponseQuotaLink, Tags", "finished/partial, JSON answers/variables/ttc/meta"],
        ["Contacts", "Contact, ContactAttributeKey/Value, Segment", "Email/userId/custom attributes, filter theo survey interaction"],
        ["Unify", "FeedbackDirectory, FeedbackSource, mappings", "Workspace association và tenant isolation"],
        ["Analysis", "Chart, Dashboard, DashboardWidget", "Widget placement, chart query và creator metadata"]
    ])
    add_callout(doc, "Multi-tenancy", "Mọi query dữ liệu nghiệp vụ phải scope Organization hoặc Workspace; Feedback Hub dùng FeedbackDirectory tenant. Không chấp nhận truy vấn chỉ bằng ID resource nếu không kiểm ownership/access.", "risk")
    source(doc, "packages/database/schema.prisma; packages/database/migration/*")

    heading(doc, "6. Danh mục chức năng nghiệp vụ")
    features = [
        ["AI LAB Survey / Import", "CSV/XLSX, preview, diagnostics, commit draft, registry/version", "Đã có; production E2E"],
        ["Khảo sát", "Tạo/sửa/duplicate/publish/pause/archive, link/app survey", "Core production"],
        ["Question types", "Open text, single/multi, NPS, rating, CSAT, CES, matrix, ranking, date, consent, contact, file", "Core + importer mapping"],
        ["Logic/adaptive", "Condition, branch, variables, calculation, quota, ending", "Compiler/runtime; cần regression theo testbank"],
        ["Danh bạ", "Contact, attributes, CSV, personal links, bulk API", "Core production"],
        ["Segment", "Dynamic filters, attributes, survey interaction, date", "Core; phụ thuộc có contact data"],
        ["Unify sources", "Nguồn Formbricks/API và mapping fields", "Beta production"],
        ["Feedback records", "List/filter/detail/manual record, sentiment/emotion", "Beta; AI coverage không đồng đều"],
        ["Taxonomy", "Run, topics/subtopics, rename/remove, record counts", "AI clustering thật; labels curated"],
        ["Charts", "Manual metrics/dimensions/filter/time/chart type", "Production"],
        ["AI chart", "Natural-language query → structured chart query", "Đã sửa schema filter và deploy"],
        ["Dashboards", "Create, widget grid, duplicate/add chart", "Production; CUID2 integrity đã sửa"],
        ["Settings workspace", "General, teams, languages, app connection, integrations, look, user actions, tags", "Production"],
        ["Settings organization", "General, teams, API keys, datasets, domain, enterprise", "Production/enterprise gated"],
        ["Account", "Profile, notifications, authorized apps, 2FA", "Production"],
        ["SSO", "Google/GitHub/Azure/OIDC/SAML + recovery/provisioning", "Code complete; UAT per IdP"],
        ["API/Webhook", "Management API, client API, v3, webhook/integrations", "Production; scope keys"],
        ["i18n/branding", "Vietnamese default, AILAB logo/name/email branding", "Phần lớn đạt; runtime validation còn English"],
    ]
    add_table(doc, ["Module", "Khả năng", "Trạng thái bàn giao"], features)

    heading(doc, "7. Importer CSV/XLSX — module mở rộng AILAB")
    add_figure(doc, diagrams["import"], "Hình 5 — BPMN rút gọn của importer")
    doc.add_paragraph("Importer là module mới nằm tại `apps/web/modules/ai-lab-survey`, không chỉ là upload CSV. Module thực hiện file guard, parse/normalize locale, canonical validation, preview diagnostics, quyết định commit và ghi transaction vào Survey cùng AILabSurveyRegistry/Version/ImportJob. Thiết kế này tách validation khỏi mutation, cho phép dry-run và audit.")
    add_table(doc, ["Lớp", "Quy ước"], [
        ["File", "CSV hoặc XLSX, tối đa 10 MiB; UTF-8 và fallback windows-1258"],
        ["CSV class", "S/SL/G/Q/V/E/EQ/CALC/A/SQ/R đại diện survey, language, group, question, variable, equation, calculation, answer, subquestion, routing"],
        ["XLSX sheets", "Survey, Groups, Questions, Options, Logic, Variables, Quotas; Guide/DataDictionary/ExpressionExamples/Compatibility tùy chọn"],
        ["Question type", "openText, multipleChoiceSingle/Multi, nps, rating, csat, ces, matrix, ranking, date"],
        ["Logic", "Comparison, boolean AND/OR/NOT, arithmetic, relevance/branch, score/variable"],
        ["Commit", "Chỉ mutate sau validated preview; transaction tạo draft, registry, version và import job"],
        ["Idempotency", "Hash/source metadata và lifecycle dùng để chống commit nhầm/đúp; lỗi trả diagnostics"],
    ])
    heading(doc, "7.1 Quy trình người dùng", 2)
    add_numbered(doc, [
        "Tải template chuẩn từ màn hình Import CSV/XLSX; không đổi tên sheet/cột bắt buộc.",
        "Điền metadata Survey, group/block, question, options, logic, variables và quotas.",
        "Upload tệp; đọc toàn bộ diagnostics theo sheet/row/field trước khi tiếp tục.",
        "Kiểm preview canonical: số block, question, option, branch, locale và ending.",
        "Chọn tạo Formbricks draft; mở editor kiểm tra logic và giao diện desktop/mobile.",
        "Publish sau khi chạy route cases, boundary cases và UAT; giữ file nguồn và version hash để audit."
    ])
    heading(doc, "7.2 Skill cho ChatGPT/AI agent khác", 2)
    doc.add_paragraph("Skill `form-builder` cung cấp contract để agent khác sinh CSV/XLSX phù hợp importer, gồm mapping LimeSurvey ExpressionScript, adaptive routing và testbank. Khi chuyển skill sang agent khác, phải kèm template chuẩn, yêu cầu giữ stable codes và checklist validate; agent không được tự phát minh cột.")
    source(doc, "apps/web/modules/ai-lab-survey/*; packages/survey-compiler/*; .agents/skills/form-builder/SKILL.md; apps/web/public/sample-csv/*")

    heading(doc, "8. Adaptive survey và Expression routing")
    add_figure(doc, diagrams["adaptive"], "Hình 6 — Sequence runtime adaptive")
    add_table(doc, ["Khái niệm", "Diễn giải", "Ví dụ"], [
        ["Stable code", "Mã câu hỏi/biến không đổi qua version", "CONSENT, ROLE, SCORE_AI"],
        ["Relevance", "Điều kiện quyết định block/question được hiển thị", "CONSENT == 'Y' && ROLE == 'STUDENT'"],
        ["Branch", "Đi đến block/ending theo đáp án/điểm", "CONSENT=N → decline ending"],
        ["Variable", "Giữ state trung gian dùng cho logic/report", "TOTAL = Q1 + Q2 + Q3"],
        ["Calculation", "Tính điểm/level có kiểm kiểu và missing", "if(TOTAL >= 80, 'L4', 'L2')"],
        ["Quota", "Dừng/redirect/mark khi đạt ngưỡng", "ROLE=MANAGER cap 100"],
        ["Adaptive path", "Chỉ hiển thị ngân hàng câu hỏi phù hợp profile/level", "8 banks theo nhóm đối tượng"],
    ])
    add_callout(doc, "Kết quả Playwright", "Khảo sát adaptive render được desktop/mobile. Case CONSENT=N của smoke spec hiện chọn sai vì test giả định nút Next/câu đầu cũ; production survey đang vào block profile và còn nhãn validation English. Cần cập nhật selector/fixture theo stable code thay vì text/order.", "warn")
    source(doc, "packages/survey-compiler; packages/surveys; apps/web/playwright/production/full-testbank-*; Testbank/*")

    heading(doc, "9. AI-native — kiến trúc và tiêu chí 'không mock'")
    add_figure(doc, diagrams["ai"], "Hình 7 — Pipeline AI feedback end-to-end")
    heading(doc, "9.1 Định nghĩa acceptance", 2)
    add_bullets(doc, [
        "Request đi qua provider/model thật và có bằng chứng network/job log; không dùng fixture response.",
        "Output được ghi vào database production đúng tenant và có timestamp/model metadata khi hỗ trợ.",
        "Đọc lại qua API/UI cho kết quả nhất quán; duplicate/idempotency không tạo record kép.",
        "Missing auth bị 401/403; tenant A không đọc được tenant B.",
        "Failure có retry/error state, không treo vô hạn ở 0/N mà không cảnh báo."
    ])
    heading(doc, "9.2 Bằng chứng đã xác minh", 2)
    add_table(doc, ["Test", "Kết quả", "Ý nghĩa"], [
        ["Embedding canary", "gemini-embedding-001, vector 768 chiều", "Model thật đã chạy và data được lưu"],
        ["Semantic search", "Canary trả về score 0.74035", "Vector được truy vấn, không phải danh sách mock"],
        ["Sentiment", "very_positive, score 1", "Gemini 2.5 Flash phân tích nội dung canary"],
        ["Emotion", "joy", "Emotion label được lưu/đọc lại"],
        ["Tenant isolation", "Passed", "Cross-tenant access bị ngăn"],
        ["Missing auth", "401", "Hub không public dữ liệu"],
        ["Duplicate insertion", "409; row count=1", "Idempotency hoạt động"],
        ["Batch corpus", "815 embedded; 280 text sentiments; 171 emotions", "Pipeline có độ phủ thật nhưng không đồng nhất"],
        ["Taxonomy", "13 clusters, 36 nodes", "Clustering thật; nhãn current đã curated"],
    ])
    heading(doc, "9.3 AI chart", 2)
    doc.add_paragraph("AI chart nhận prompt tự nhiên, cung cấp context fields/metrics cho model, yêu cầu structured output rồi normalize/validate thành chart query. Lỗi production xảy ra khi model trả filter operator `set` kèm values; schema cũ reject. Bản sửa bỏ values thừa cho `set/notSet`, giữ strict validation cho operator cần value và đã có regression test tái hiện exact output.")
    source(doc, "apps/web/modules/ee/analysis/charts/lib/ai-chart-query.server.ts; ai-chart-query.test.ts; apps/web/modules/hub/*; apps/web/app/api/v3/unify-feedback/*")

    heading(doc, "10. Dashboard, chart và khai phá điểm số")
    add_table(doc, ["Nhóm chỉ số", "Công thức/nguồn", "Visual phù hợp", "Lưu ý"], [
        ["NPS điểm số", "(% Promoters - % Detractors)", "KPI, trend, breakdown", "Chỉ survey có NPS"],
        ["CSAT điểm số", "% rating 4–5 trên thang 1–5", "KPI/bar/time", "Cần đúng question type"],
        ["CES điểm số", "Theo scale cấu hình", "KPI/distribution", "Không trộn scale khác"],
        ["Average NPS/CSAT", "Average raw rating", "Line/bar", "Khác score chuẩn hóa"],
        ["Response count", "count responses", "KPI/time", "Phân biệt finished/partial"],
        ["Sentiment/emotion", "Hub AI fields", "Stacked bar/heatmap", "Missing label phải hiện Unknown"],
        ["Topic/subtopic", "Taxonomy assignment", "Treemap/bar/drilldown", "Label quality cần review"],
        ["Adaptive level", "Variables/calculated output", "Funnel/cohort", "Cần expose field trong semantic model"],
    ])
    add_callout(doc, "Lỗi 'Tài nguyên không tồn tại'", "Dashboard/widget phải tham chiếu chart ID CUID2 hợp lệ, đúng workspace và còn tồn tại. Seed/demo từng dùng ID sai format dẫn đến trang resource error; dữ liệu hiện đã được sửa, nhưng migration/seed nên validate FK và ownership trước deploy.", "risk")

    heading(doc, "11. Danh bạ, thuộc tính và segment")
    add_table(doc, ["Use case", "Luồng", "Acceptance"], [
        ["Import contacts", "CSV → detect data type → validate email/attributes → bulk create", "Báo row lỗi; không mất dòng tốt; tenant scope"],
        ["Custom attributes", "Tạo key → nhập/update values", "Đúng string/number/date/boolean; unique policy"],
        ["Dynamic segment", "Filter attributes + survey interaction + date", "Preview count khớp Prisma query; save/load ổn định"],
        ["Personal link", "Contact/segment → survey link", "Mỗi contact đúng token/link; revoke/expiry theo policy"],
        ["Empty state", "Không có contact", "UI giải thích và CTA upload CSV; không coi empty là lỗi chức năng"],
    ])
    source(doc, "apps/web/modules/ee/contacts; apps/web/playwright/segments.spec.ts; apps/web/app/api/v2/management/contacts")

    heading(doc, "12. Unify Feedback, sentiment/emotion và taxonomy")
    add_table(doc, ["Giai đoạn", "Input", "Output", "Failure mode / xử lý"], [
        ["Source mapping", "Survey/API fields", "Canonical feedback record", "Mapping thiếu → diagnostics"],
        ["Embedding", "Open text", "Vector + progress", "Job queue/credential/model limit"],
        ["Sentiment/emotion", "Text đủ nghĩa", "Labels + scores", "Text rỗng/ngắn → Unknown hợp lệ"],
        ["Taxonomy run", "Embedded corpus", "Cluster/node tree/counts", "Model label rác → human review/rename"],
        ["Drill-down", "Topic/subtopic", "Record list", "Access scope và pagination"],
    ])
    add_callout(doc, "0 / 815 đứng lâu", "UI phải hiển thị job state, last heartbeat, retry và lỗi cấu hình provider/worker; không chỉ polling count. Checklist vận hành ở Chương 17 yêu cầu kiểm worker log, queue lag, Hub API key, model credential, DB write và tenant ID.", "warn")

    heading(doc, "13. Identity, SSO, RBAC và bảo mật")
    add_figure(doc, diagrams["rbac"], "Hình 8 — Use case RBAC và tenant boundary")
    add_table(doc, ["Cấp", "Vai trò/quyền", "Phạm vi"], [
        ["Organization", "owner, manager, member", "Billing, domain, enterprise, teams, API keys"],
        ["Team", "admin, contributor", "Nhóm người dùng trong organization"],
        ["WorkspaceTeam", "read, readWrite, manage", "Workspace và resource bên trong"],
        ["API key", "Permissions + organizationAccess + workspace links", "Management API theo scope"],
        ["SSO", "SAML/OIDC/OAuth provisioning/recovery", "Account linking và org/team mapping"],
    ])
    heading(doc, "13.1 Security controls cần duy trì", 2)
    add_bullets(doc, [
        "Không ghi token invite, session, API key, SMTP/OAuth secret vào tài liệu, log hoặc screenshot.",
        "Mọi server action/route/service kiểm membership và workspace ownership; UI guard chỉ là lớp bổ sung.",
        "API key lưu lookup hash, có expiry/last used và revoke; quyền tối thiểu.",
        "SSO account linking chống account takeover; recovery flow phải có audit và giới hạn quyền.",
        "2FA lockout, password reset và verification token có expiry/rate-limit.",
        "Custom scripts có browser access đầy đủ: chỉ trusted admins và cần CSP/governance."
    ])
    source(doc, "apps/web/modules/ee/sso; apps/web/modules/ee/auth/saml; packages/database/schema.prisma; apps/web/lib/membership")

    heading(doc, "14. i18n, thương hiệu và email")
    add_table(doc, ["Hạng mục", "Chuẩn yêu cầu", "Hiện trạng / hành động"], [
        ["App locale", "Tiếng Việt mặc định theo user/app source of truth", "Shell/menu phần lớn Việt hóa; kiểm locale cookie/profile"],
        ["Survey runtime", "Button, required, validation, accessibility label đều Việt", "Còn Required, Please fill out this field, Rate X out of Y"],
        ["Date/time", "Shared formatter theo app locale", "Không dùng browser default/toLocale ad hoc"],
        ["Brand", "AILAB Survey logo/name/color", "Logo production đã đổi; kiểm tất cả auth/email/error pages"],
        ["Invite email", "AILAB Survey, tiếng Việt, MAIL_FROM uelailab@gmail.com", "Cần end-to-end SMTP inbox test sau mỗi template/deploy"],
        ["Source translation", "Chỉ sửa en-US rồi chạy pnpm i18n", "Không hand-edit locale generated theo repo policy"],
    ])
    if (ASSET_DIR / "production-login-vietnamese.png").exists():
        add_figure(doc, ASSET_DIR / "production-login-vietnamese.png", "Ảnh production — trang đăng nhập đã dùng nhận diện AILAB Survey và tiếng Việt", width=6.5)
    if (ASSET_DIR / "production-adaptive-validation.png").exists():
        add_figure(doc, ASSET_DIR / "production-adaptive-validation.png", "Ảnh production — survey adaptive còn chuỗi validation/rating tiếng Anh cần đóng i18n", width=6.5)
    source(doc, "apps/web/locales/en-US.json; packages/surveys/src/locales; packages/email; Playwright production trace 04/08/2026")

    heading(doc, "15. Hướng dẫn thao tác theo vai trò")
    role_guides = {
        "End user — trả lời khảo sát": ["Mở link", "Chọn ngôn ngữ", "Đọc consent", "Trả lời theo block thích ứng", "Xác nhận kết quả/ending"],
        "Survey designer": ["Tạo/import draft", "Thiết kế block/question", "Cấu hình logic/variables/quota", "Preview desktop/mobile", "Publish và theo dõi responses"],
        "Data analyst": ["Kiểm nguồn Unify", "Theo dõi embedding", "Review sentiment/taxonomy", "Tạo chart", "Gắn dashboard và drill-down"],
        "Workspace admin": ["Quản lý teams/languages/look", "Kết nối app/integration", "Tags/user actions", "API key theo scope", "Audit và cleanup"],
        "Organization admin": ["Membership/roles", "Team mapping", "Domain/SSO", "Dataset", "Enterprise license/billing"],
        "Operator/DevOps": ["Health/metrics", "DB migration", "Worker queues", "Backup/restore", "Incident rollback"],
    }
    for role, steps in role_guides.items():
        heading(doc, role, 2)
        add_numbered(doc, steps)

    heading(doc, "15.1 Tài khoản demo production", 2)
    add_callout(doc, "Thông tin nhạy cảm", "Tài khoản dưới đây chỉ dùng cho test/demo có kiểm soát. Mật khẩu là mật khẩu tạm thời; phải đổi ngay sau buổi demo hoặc vô hiệu hóa/xóa tài khoản. Không gửi tài liệu này ra ngoài phạm vi được ủy quyền.", "risk")
    add_table(doc, ["Thuộc tính", "Giá trị"], [
        ["URL đăng nhập", "https://formailab.royalai.dev/auth/login"],
        ["Email", "ailab.demo.test@royalai.dev"],
        ["Mật khẩu tạm thời", demo_password],
        ["Workspace", "My workspace — cms7eskq6000901o9ul1ipheu"],
        ["Organization role", "member"],
        ["Team", "AILAB Demo Testers — contributor"],
        ["Workspace permission", "readWrite: xem, tạo, cập nhật; không xóa"],
        ["Locale", "vi-VN"],
        ["Kiểm chứng", "Playwright production: login HTTP 200, vào trang Surveys, thấy heading Khảo sát"],
        ["Hạn chế đã xác nhận", "Không có owner/manager; không quản trị billing, domain, enterprise, API key hoặc xóa tài nguyên"],
    ], widths=[5, 11], font_size=9)
    add_callout(doc, "Phát hiện i18n khi xác minh", "Nội dung trang và heading đã là tiếng Việt nhưng HTML title vẫn là `Your Surveys | AILAB Survey`. Đưa title này vào backlog Việt hóa P0/P1.", "warn")

    heading(doc, "15.2 Bản đồ điều hướng cho end user", 2)
    add_table(doc, ["Khu vực", "Mục đích", "Khi nào sử dụng"], [
        ["Import CSV / XLSX", "Biến ngân hàng câu hỏi/template thành survey draft", "Khởi tạo survey lớn, adaptive hoặc migrate LimeSurvey"],
        ["Khảo sát", "Quản lý vòng đời survey và phản hồi", "Tạo, sửa, publish, chia sẻ, xem kết quả"],
        ["Danh bạ", "Quản lý contact, thuộc tính và phân đoạn", "Targeting, personal link, khảo sát theo nhóm"],
        ["Dữ liệu phản hồi", "Hợp nhất nguồn feedback có cấu trúc và open text", "Sentiment, emotion, taxonomy, semantic analysis"],
        ["Phân tích", "Tạo chart và dashboard", "NPS/CSAT/CES, trend, breakdown và AI chart"],
        ["Cài đặt", "Cấu hình workspace, teams, language, integration, look", "Thiết lập hệ thống và quyền truy cập"],
        ["Tổ chức", "Quản lý members, teams, API keys, datasets, domain, license", "Owner/manager; demo member không được dùng"],
        ["Tài khoản", "Profile, notification, authorized apps", "Thiết lập cá nhân và bảo mật"],
    ])

    heading(doc, "15.3 Import CSV/XLSX — hướng dẫn chi tiết", 2)
    add_table(doc, ["Bước", "Thao tác end user", "Hệ thống thực hiện", "Kết quả/kiểm tra"], [
        ["1", "Tải template CSV/XLSX chuẩn", "Cung cấp đúng schema/sheet/cột", "Giữ nguyên tên sheet và stable codes"],
        ["2", "Điền Survey, Groups, Questions, Options", "Parser ánh xạ sang canonical survey", "Đủ tiêu đề, type, required, scale/options"],
        ["3", "Điền Logic, Variables, Quotas", "Compiler kiểm expression và reference", "Không tham chiếu code không tồn tại"],
        ["4", "Upload tệp ≤10 MiB", "File guard kiểm MIME/size/encoding", "Tệp được nhận hoặc lỗi rõ nguyên nhân"],
        ["5", "Đọc preview/diagnostics", "Validate toàn file, không mutate DB", "Sửa tất cả error; warning phải được chấp thuận"],
        ["6", "Chọn tạo draft", "Transaction tạo Survey + registry/version/job", "Chuyển tới editor đúng survey mới"],
        ["7", "Preview và chạy route cases", "Runtime biên dịch branch/score/ending", "Desktop/mobile và các nhánh đều đúng"],
        ["8", "Publish", "Đổi lifecycle/status và tạo public link", "Link truy cập được; version/hash được lưu"],
    ])

    heading(doc, "15.4 Khảo sát — danh sách, editor và vòng đời", 2)
    add_table(doc, ["Chức năng", "Cách dùng", "Giải thích core logic"], [
        ["Tìm kiếm/lọc/sắp xếp", "Nhập tên; chọn status/type/order", "Query workspace-scoped, phân trang/cursor; không tìm ngoài tenant"],
        ["Khảo sát mới", "Chọn tạo từ đầu, template hoặc import", "Tạo draft, creator và workspace ownership"],
        ["Editor", "Thêm block/question/ending và kéo thả", "Questions/blocks/endings lưu JSON có schema validation"],
        ["Question types", "Chọn open text, single/multi, NPS, rating, CSAT, CES, matrix, ranking, date…", "Mỗi type có response contract và metric tương ứng"],
        ["Required/validation", "Bật bắt buộc, min/max, regex/rules", "Runtime chặn next/submit và trả validation localized"],
        ["Logic", "Chọn câu hỏi + operator + value + action", "Evaluator quyết định show/hide/jump/ending theo stable IDs"],
        ["Variables/calculation", "Khai báo biến và công thức điểm", "Giữ state trong response.variables; dùng cho branch/report"],
        ["Quota", "Đặt điều kiện, giới hạn và action", "Đếm partial/finished theo cấu hình; tạo ResponseQuotaLink"],
        ["Preview", "Chạy survey không tạo response chính thức", "Kiểm UI/logic trên bundle runtime hiện hành"],
        ["Publish/pause", "Đưa survey live hoặc tạm dừng", "Status gate kiểm public access và scheduling"],
        ["Duplicate/archive", "Nhân bản hoặc lưu trữ", "Tạo ID mới; archive không xóa response lịch sử"],
        ["Summary/responses", "Xem completion, response table và export", "Đọc Response scoped survey/workspace; phân biệt finished/partial"],
    ])

    heading(doc, "15.5 Danh bạ, thuộc tính và phân đoạn", 2)
    add_table(doc, ["Chức năng", "Thao tác", "Kết quả mong đợi"], [
        ["Tải CSV contacts", "Chuẩn bị email/firstName/lastName/userId và custom fields", "Preview row errors, import bản ghi hợp lệ, không duplicate unique key"],
        ["Tạo contact", "Nhập định danh và thuộc tính", "Contact gắn đúng workspace"],
        ["Thuộc tính", "Tạo key string/number/date/boolean", "Giá trị được validate và format nhất quán"],
        ["Phân đoạn", "Ghép filter AND/OR theo attributes", "Preview count khớp records"],
        ["Survey interaction", "Lọc theo đã/chưa trả lời survey và thời gian", "Prisma query phản ánh đúng response/display"],
        ["Personal link", "Chọn contact/segment rồi sinh link", "Mỗi contact nhận link định danh đúng survey"],
        ["Empty state", "Nếu chưa có dữ liệu, upload CSV hoặc tạo contact", "Không có kết quả là trạng thái dữ liệu, không mặc định là lỗi module"],
    ])

    heading(doc, "15.6 Dữ liệu phản hồi và AI-native", 2)
    add_table(doc, ["Chức năng", "Thao tác", "Điều cần hiểu"], [
        ["Nguồn phản hồi", "Kết nối Formbricks/API và map fields", "Nguồn phải liên kết FeedbackDirectory và workspace"],
        ["Bản ghi", "Tìm/filter/mở drawer chi tiết", "Canonical record có source, createdAt và fields"],
        ["Embedding progress", "Theo dõi X/N và trạng thái job", "Worker gọi model thật và ghi vector; 0/N lâu cần runbook"],
        ["Sentiment", "Xem label/score", "Dấu — có thể là pending, not-applicable hoặc error; không tự suy diễn neutral"],
        ["Emotion", "Xem emotion labels", "Chỉ sinh khi text đủ điều kiện và job thành công"],
        ["Taxonomy", "Trigger run, mở topic/subtopic, drill-down records", "Clustering AI thật; label cần human quality gate"],
        ["Rename/remove", "Curate node taxonomy", "Thay đổi presentation/governance, không bịa lại source text"],
        ["Semantic search", "Tìm bằng ý nghĩa thay vì keyword", "So cosine/vector score trong cùng tenant"],
    ])

    heading(doc, "15.7 Biểu đồ và bảng điều khiển", 2)
    add_table(doc, ["Chức năng", "Cách dùng", "Core logic/metric"], [
        ["Manual chart", "Chọn measure, dimension, filter, time và chart type", "Sinh Cube query có schema validation"],
        ["Built-in score", "Chọn NPS/CSAT/CES score hoặc average", "Score chuẩn hóa khác raw average; chỉ hiện khi survey có type phù hợp"],
        ["AI chart", "Nhập yêu cầu tiếng Việt và review query", "LLM sinh structured query; normalize rồi validate, không chạy text SQL tùy ý"],
        ["Preview/data viewer", "Xem chart và bảng dữ liệu trước save", "Phát hiện empty/missing dimension và scale sai"],
        ["Save chart", "Đặt tên rõ metric/segment/time", "Chart thuộc workspace và creator"],
        ["Add dashboard", "Chọn dashboard và vị trí widget", "DashboardWidget giữ chartId, placement và layout"],
        ["Drill-down", "Mở điểm dữ liệu/record", "Access vẫn kiểm workspace và resource tồn tại"],
        ["Resource error", "Nếu 404/không quyền, kiểm chart/widget ID", "CUID2, FK và workspace ownership phải nhất quán"],
    ])

    heading(doc, "15.8 Cài đặt workspace và tổ chức", 2)
    add_table(doc, ["Màn hình", "Chức năng", "Quyền/lưu ý"], [
        ["Chung", "Tên workspace, cooldown, custom head scripts, ID", "Script có browser access; chỉ trusted admin"],
        ["Quyền truy cập nhóm", "Map team → read/readWrite/manage", "Member cần team để vào workspace"],
        ["Ngôn ngữ khảo sát", "Bật locale/default và language switch", "Tách locale hiển thị khỏi timezone"],
        ["Kết nối ứng dụng", "Website/app SDK setup", "Workspace ID và environment contract"],
        ["Tích hợp", "Webhook/Slack/Sheets/Airtable/Notion…", "Secret/signature và least privilege"],
        ["Giao diện", "Logo, màu, placement, branding", "Phân biệt workspace look và survey override"],
        ["Hành động người dùng", "Event/action classes cho targeting", "Dùng cho app survey triggers"],
        ["Thẻ", "Phân loại survey/response", "Workspace-scoped; hỗ trợ reporting"],
        ["Xóa workspace", "Xóa survey/response/contact/action", "Destructive và không hoàn tác; demo readWrite không được xóa"],
        ["Organization teams", "Members, roles, team mappings", "Owner/manager"],
        ["API keys", "Tạo/revoke và scope org/workspace", "Không hiển thị/lưu key trong tài liệu"],
        ["Tập dữ liệu", "FeedbackDirectory ↔ workspace", "Kiểm conflict và tenant isolation"],
        ["Domain/SSO/license", "Custom domain, IdP, enterprise entitlement", "Enterprise/owner; cần UAT riêng"],
    ])

    heading(doc, "15.9 Checklist demo bằng tài khoản test", 2)
    add_bullets(doc, [
        "Đăng nhập và xác nhận workspace selector hiển thị `My workspace`.",
        "Mở danh sách Khảo sát; tìm `AILAB 120Q Advanced Adaptive Intelligence Assessment 2026`.",
        "Tạo một survey draft tên có tiền tố `[DEMO]`; không publish nếu chưa được phê duyệt.",
        "Mở Import CSV/XLSX, upload template và dừng tại preview nếu buổi demo không cho phép mutation.",
        "Tạo contact/segment demo có prefix `[DEMO]`; không dùng dữ liệu cá nhân thật.",
        "Tạo chart/dashboard demo từ dữ liệu hiện có; không xóa chart/dashboard của người khác.",
        "Kiểm Unify progress/sentiment/taxonomy ở chế độ xem; không trigger backfill lớn trong giờ production.",
        "Kết thúc: xóa dữ liệu `[DEMO]` bằng admin được ủy quyền, đổi mật khẩu hoặc vô hiệu hóa tài khoản demo."
    ])

    heading(doc, "16. Demo kịch bản cho BA/PM/PO/Stakeholder")
    demos = [
        ["DEMO-01", "Import testbank", "Upload template 120Q → preview 112 questions/508 options → draft", "Importer không chỉ đọc file mà tạo survey thật"],
        ["DEMO-02", "Adaptive consent", "CONSENT=N → decline ending; Y → profile/bank", "Branch deterministic"],
        ["DEMO-03", "Adaptive level", "Trả lời score thấp/cao → bộ câu hỏi/level khác nhau", "Personalized assessment"],
        ["DEMO-04", "Contacts/segment", "Import contacts → attribute → segment → personal link", "Targeting end-to-end"],
        ["DEMO-05", "AI feedback", "Ingest open text → embedding/sentiment/emotion", "Model thật + persisted data"],
        ["DEMO-06", "Taxonomy", "Run 815 records → topics/subtopics → drill down", "Unstructured feedback discovery"],
        ["DEMO-07", "AI chart", "Prompt tiếng Việt → chart query → save dashboard", "Natural-language analytics"],
        ["DEMO-08", "RBAC/SSO", "Member vs workspace manager; SSO provision", "Tenant/security boundary"],
        ["DEMO-09", "Localization/brand", "Login, email invite, survey runtime", "AILAB identity nhất quán"],
        ["DEMO-10", "Operations", "Health, logs, job progress, retry/rollback", "Production observability"],
    ]
    add_table(doc, ["ID", "Kịch bản", "Thao tác", "Thông điệp"], demos)
    if (ASSET_DIR / "production-adaptive-survey.png").exists():
        add_figure(doc, ASSET_DIR / "production-adaptive-survey.png", "Ảnh production — khảo sát adaptive render trên trình duyệt thực", width=6.5)

    heading(doc, "17. Chiến lược kiểm thử và tiêu chí nghiệm thu")
    add_table(doc, ["Lớp test", "Phạm vi", "Gate"], [
        ["Unit/Vitest", "Parser, validator, expression, transformer, AI schema, RBAC utilities", "Không regression; deterministic"],
        ["Integration", "Prisma transaction, importer registry/version, SSO provisioning, Hub gateway", "Database thật, cleanup"],
        ["API", "Auth, tenant isolation, idempotency, pagination, error contract", "401/403/404/409 đúng"],
        ["Playwright", "Critical user journeys, adaptive routes, importer, contacts, dashboard", "Desktop + mobile, trace on failure"],
        ["AI eval", "Golden set tiếng Việt, sentiment/emotion/taxonomy/AI chart", "Quality threshold + human review"],
        ["Performance", "Import large file, 815+ embeddings, chart query, pagination", "SLO, no unbounded query"],
        ["Security", "RBAC matrix, SSO linking, API key scope, custom script", "No cross-tenant/access escalation"],
        ["DR", "Backup restore, migration rollback, worker retry", "RTO/RPO evidence"],
    ])
    heading(doc, "17.1 Production smoke 04/08/2026", 2)
    add_table(doc, ["Case", "Kết quả", "Phân tích"], [
        ["Health dependencies", "PASS", "Web, main database và cache phản hồi tốt"],
        ["Auth guard", "FAIL assertion", "Redirect/login đúng; test kỳ vọng title English trong khi production đã Vietnamese"],
        ["Excel template download", "PASS", "HTTP 200, PK/XLSX và kích thước hợp lệ"],
        ["Adaptive desktop/mobile", "PASS", "Survey inProgress render trên Chrome"],
        ["CONSENT=N smoke", "FAIL journey", "Spec click Next theo survey cũ; survey hiện bắt đầu bằng profile block và validation English"],
    ])
    add_callout(doc, "Không đánh đồng test fail với outage", "Hai failure cho thấy test contract/fixture đã drift và i18n runtime chưa đủ; health/template/render vẫn pass. Cần sửa smoke test dùng stable codes/data-testid và một survey fixture versioned.", "info")
    source(doc, "test-results/production/junit.xml; test-results/production/results.json; trace.zip/video/screenshots")

    heading(doc, "18. Vận hành, quan sát và xử lý sự cố")
    runbooks = [
        ["Embedding 0/N", "Hub worker không nhận job, provider key/model, queue/DB write", "Kiểm worker log → queue lag → tenant/source ID → model call → record state → retry có giới hạn"],
        ["Chart resource error", "ID/FK/workspace ownership hoặc chart bị xóa", "Đối chiếu DashboardWidget.chartId, CUID2, workspaceId; repair FK/seed"],
        ["AI chart invalid", "Model output ngoài schema", "Lưu sanitized output, normalize valueless operators, regression fixture, fallback manual builder"],
        ["Sentiment/emotion —", "Text rỗng/ngắn, job chưa chạy hoặc model fail", "Phân biệt not-applicable/pending/error; retry only eligible records"],
        ["Email còn Formbricks", "Template/env/image cache hoặc service chưa redeploy", "Kiểm MAIL_FROM/NAME, template render, asset URL, SMTP inbox, restart"],
        ["UI còn English", "Key thiếu ở source/survey bundle cache", "Sửa en-US → pnpm i18n → force survey build → deploy → hard refresh"],
        ["SSO fail", "Metadata/cert/redirect/domain/team mapping", "Validate IdP config, callback, clock skew, recovery path; không auto-link mù"],
    ]
    add_table(doc, ["Triệu chứng", "Nguyên nhân thường gặp", "Runbook"], runbooks)
    heading(doc, "18.1 SLO/monitoring đề xuất", 2)
    add_table(doc, ["Chỉ số", "Mục tiêu đề xuất", "Alert"], [
        ["Web availability", "≥99.9%", "5xx/health fail 3 phút"],
        ["API p95", "<800 ms core; <3 s analytics", "p95 vượt 2 cửa sổ"],
        ["Embedding freshness", "95% <10 phút", "Queue oldest >15 phút"],
        ["AI error rate", "<2% không tính invalid input", ">5%/10 phút"],
        ["DB connection/index", "No pool saturation", ">80% pool hoặc slow query"],
        ["Backup", "Daily + restore drill quarterly", "Missed backup/restore checksum"],
    ])

    heading(doc, "19. Rủi ro, nợ kỹ thuật và backlog ưu tiên")
    risks = [
        ["P0", "Runtime survey còn chuỗi English", "Ảnh hưởng trải nghiệm tiếng Việt mặc định", "Bổ sung locale keys/validation/rating aria; force rebuild survey bundle; E2E vi"],
        ["P0", "Embedding progress có thể treo không rõ lỗi", "Người dùng không biết job failure", "Persist state/heartbeat/error/retry; operator dashboard"],
        ["P1", "Taxonomy label model tạo gibberish", "Chủ đề không sử dụng được", "Constrained prompt/schema, language detect, quality gate, human approval"],
        ["P1", "Smoke tests drift theo text/order", "False negative và bỏ sót route", "Stable codes/data-testid, versioned fixture"],
        ["P1", "License API 400 fallback", "Feature entitlement có thể không nhất quán", "Sửa license config/contract, alert; không chỉ DB toggle"],
        ["P1", "MaxListenersExceededWarning", "Rò listener hoặc false warning", "Trace listener lifecycle, cleanup, load test"],
        ["P1", "SSO chưa UAT mọi IdP", "Rủi ro provisioning/account linking", "Matrix Google/GitHub/Azure/OIDC/SAML + recovery"],
        ["P2", "Demo objects tên xq/xz/a", "Khó bàn giao và quản trị", "Archive/rename demo, seed catalog chuẩn"],
        ["P2", "AI coverage 280/815 sentiment", "Dashboard có nhiều Unknown", "Eligibility rules, backfill, progress by stage"],
    ]
    add_table(doc, ["Ưu tiên", "Rủi ro", "Tác động", "Hành động/DoD"], risks)

    heading(doc, "20. Checklist bàn giao và go-live")
    checks = [
        "[ ] Chốt commit/image digest và lưu release notes.",
        "[ ] Migration deploy chạy thành công; có backup và restore test.",
        "[ ] Health web/database/cache/Hub/worker/taxonomy/Cube đều xanh.",
        "[ ] Import CSV/XLSX dry-run + commit + cleanup đạt trên production-isolated tenant.",
        "[ ] Adaptive route golden cases, boundary cases và mobile đạt.",
        "[ ] Embedding/sentiment/emotion/taxonomy canary thật được ghi và đọc lại.",
        "[ ] Tenant isolation, auth 401, duplicate 409 và API key scope đạt.",
        "[ ] AI chart prompt tiếng Việt + manual chart + dashboard drilldown đạt.",
        "[ ] Email invite nhận đúng AILAB logo/name/from và link hợp lệ; không lộ token trong ticket/doc.",
        "[ ] Vietnamese scan không còn English P0 trong app/survey runtime.",
        "[ ] UAT role matrix org/team/workspace và SSO IdP thực tế đạt.",
        "[ ] SLO dashboard, alerts, on-call, rollback và owner được phê duyệt."
    ]
    add_bullets(doc, checks)

    heading(doc, "21. Phụ lục A — Route và màn hình trọng yếu")
    routes = [
        ["Import CSV/XLSX", "/workspaces/{workspaceId}/ai-lab-survey"],
        ["Surveys", "/workspaces/{workspaceId}/surveys"],
        ["Contacts", "/workspaces/{workspaceId}/contacts"],
        ["Segments", "/workspaces/{workspaceId}/segments"],
        ["Unify sources", "/workspaces/{workspaceId}/unify/sources"],
        ["Feedback records", "/workspaces/{workspaceId}/unify/feedback-records"],
        ["Taxonomy", "/workspaces/{workspaceId}/unify/taxonomy"],
        ["Dashboards", "/workspaces/{workspaceId}/dashboards"],
        ["Charts", "/workspaces/{workspaceId}/charts"],
        ["Workspace settings", "/workspaces/{workspaceId}/settings/workspace/general"],
        ["Organization settings", "/organizations/{organizationId}/settings/general"],
        ["Account", "/account/settings/profile"],
    ]
    add_table(doc, ["Màn hình", "Route"], routes)
    p = doc.add_paragraph("Production tham chiếu: ")
    add_hyperlink(p, "formailab.royalai.dev", "https://formailab.royalai.dev")

    heading(doc, "22. Phụ lục B — API/integration inventory")
    add_table(doc, ["Bề mặt", "Mục đích", "Auth"], [
        ["/api/v1, /api/v2 management", "Survey/contact/response management", "API key/session theo endpoint"],
        ["/api/v1, /api/v2 client", "Survey display và response từ client", "Workspace/client context"],
        ["/api/v3/unify-feedback", "Taxonomy/source/feedback orchestration", "Session + access"],
        ["/v1/feedback-records", "Proxy/gateway Hub", "Hub/API key và tenant"],
        ["Webhooks/integrations", "Slack, Google Sheets, Airtable, Notion, n8n/Activepieces", "Secret/signature/config"],
        ["MCP/OAuth provider", "Authorized apps và external agents", "OAuth consent/scopes"],
    ])

    heading(doc, "23. Phụ lục C — Traceability nguồn mã")
    sources = [
        ["AI LAB importer", "apps/web/modules/ai-lab-survey; packages/survey-compiler"],
        ["Survey runtime", "packages/surveys; packages/survey-ui; packages/js-core"],
        ["Database", "packages/database/schema.prisma; packages/database/migration"],
        ["Contacts/segments", "apps/web/modules/ee/contacts"],
        ["Charts/dashboards", "apps/web/modules/ee/analysis"],
        ["Unify/taxonomy", "apps/web/modules/ee/unify-feedback; apps/web/app/api/v3/unify-feedback"],
        ["SSO/SAML", "apps/web/modules/ee/sso; apps/web/modules/ee/auth/saml"],
        ["Email/branding", "packages/email; apps/web/public/images; env runtime"],
        ["Deployment", "docker; docker/ai-lab; docker-compose.dev.yml"],
        ["E2E evidence", "apps/web/playwright/production; test-results/production"],
        ["Agent skill", ".agents/skills/form-builder/SKILL.md"],
        ["Knowledge graph", ".ua/knowledge-graph.json; .ua/intermediate"],
    ]
    add_table(doc, ["Miền", "Nguồn chính"], sources)

    heading(doc, "24. Phụ lục D — Thuật ngữ")
    add_table(doc, ["Thuật ngữ", "Định nghĩa"], [
        ["AI-native", "AI nằm trong luồng dữ liệu/nghiệp vụ thật, có persistence, access control, retry và observability; không phải UI demo gọi mock."],
        ["Canonical survey", "Mô hình trung gian chuẩn hóa từ CSV/XLSX trước khi commit thành Survey."],
        ["Adaptive", "Luồng câu hỏi thay đổi theo profile, đáp án, biến, điểm và quota."],
        ["FeedbackDirectory", "Tenant/dataset logic của Hub liên kết với workspace."],
        ["Embedding", "Vector biểu diễn ngữ nghĩa dùng semantic search/clustering."],
        ["Taxonomy", "Cây topic/subtopic gắn các feedback records."],
        ["CUID2", "Định dạng ID dùng cho một số resource mới như chart/dashboard/demo seed."],
        ["RTO/RPO", "Mục tiêu thời gian phục hồi / điểm phục hồi dữ liệu."],
    ])

    heading(doc, "25. Kết luận bàn giao")
    doc.add_paragraph("Hệ thống đã có nền tảng kỹ thuật và nghiệp vụ rộng: survey authoring/runtime, importer adaptive mới, contacts/segments, Unify, analytics, RBAC/SSO và AI pipeline thật. Production evidence xác nhận web/data/cache, template, survey rendering và AI canary hoạt động; đồng thời phát hiện các khoảng trống quan trọng về Việt hóa runtime, job observability, taxonomy label quality và test drift. Bản tài liệu này là baseline để BA/PM/PO thống nhất phạm vi, Dev/DevOps vận hành, QA xây regression và stakeholder đánh giá go-live theo checklist thay vì theo cảm nhận giao diện.")
    add_callout(doc, "Owner action", "Đóng các mục P0 trước tuyên bố hoàn thiện toàn bộ; sau đó chạy lại full Playwright + AI eval + SSO UAT, cập nhật mục 17/19/20 và phát hành tài liệu v1.1.", "ok")

    doc.core_properties.title = "AILAB Survey — Tài liệu bàn giao kỹ thuật và nghiệp vụ"
    doc.core_properties.subject = "Full system handover: architecture, business, AI-native, operations and user guide"
    doc.core_properties.author = "UEL AI Lab"
    doc.core_properties.keywords = "AILAB Survey, Formbricks, AI-native, CSV, XLSX, adaptive survey, handover"
    doc.core_properties.comments = "Generated from repository and production evidence; contains no secrets."
    doc.core_properties.created = datetime.now(timezone.utc)
    doc.save(OUT_FILE)
    return OUT_FILE


if __name__ == "__main__":
    path = build_document()
    print(path)
